from docx import Document
from docx.oxml.ns import qn
import re

PATH = r"C:\Users\Joyce_SUN\Desktop\双引擎柔性车间调度专利书.docx"
doc = Document(PATH)

MATH_SYMBOLS = set("∈ℕℤℝ⌈⌉⌊⌋Δ·×÷√∑∏≥≤≈≠∞βαγθλμστωπ∂∫∧∨⊕⊗→⇌³²¹₀₁₂₃₄₅₆₇₈₉")
OMML_NS = '{http://schemas.openxmlformats.org/officeDocument/2006/math}'

def has_math(t):
    return any(ch in MATH_SYMBOLS for ch in t)

print("=== 含数学符号的段落 ===")
count = 0
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if has_math(t):
        count += 1
        print(f"\n[{i}] {t}")

print(f"\n=== 含数学符号的段落数: {count} ===")

# 检查 OMML 公式对象
print("\n=== OMML 公式对象 (m:oMath) ===")
omml_count = 0
for i, p in enumerate(doc.paragraphs):
    om = p._p.findall('.//' + qn('m:oMath'))
    if om:
        omml_count += len(om)
        for o in om:
            xml = o.getroottree().text if False else None
        print(f"[{i}] 段落内含 {len(om)} 个 OMML 公式")
print(f"OMML 公式总数: {omml_count}")
